import re
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def parse_tagged_resume(resume_text: str) -> list:
    segments = []
    tag_pattern = re.compile(
        r"\[(NAME|CONTACT|HEADER|BODY|BULLET)\](.*?)\[/\1\]",
        re.DOTALL,
    )

    for match in tag_pattern.finditer(resume_text):
        tag_type = match.group(1).lower()
        text = match.group(2).strip()
        if text:
            segments.append({"type": tag_type, "text": text})

    return segments


def _apply_run_style(run, style_dict):
    if style_dict.get("font_name"):
        run.font.name = style_dict["font_name"]
    if style_dict.get("font_size_pt"):
        run.font.size = Pt(style_dict["font_size_pt"])
    if style_dict.get("bold") is not None:
        run.font.bold = style_dict["bold"]
    if style_dict.get("color_hex"):
        hex_color = style_dict["color_hex"].lstrip("#")
        run.font.color.rgb = RGBColor(
            int(hex_color[0:2], 16),
            int(hex_color[2:4], 16),
            int(hex_color[4:6], 16),
        )


def _add_bottom_border(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/>'
        "</w:pBdr>"
    )
    pPr.append(pBdr)


def _set_page_margins(section, page_styles):
    if page_styles.get("margin_top_dxa"):
        section.top_margin = Twips(page_styles["margin_top_dxa"])
    if page_styles.get("margin_bottom_dxa"):
        section.bottom_margin = Twips(page_styles["margin_bottom_dxa"])
    if page_styles.get("margin_left_dxa"):
        section.left_margin = Twips(page_styles["margin_left_dxa"])
    if page_styles.get("margin_right_dxa"):
        section.right_margin = Twips(page_styles["margin_right_dxa"])


def _add_segment_to_doc(doc, segment, styles):
    seg_type = segment["type"]
    text = segment["text"]

    if seg_type == "name":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        _apply_run_style(run, styles.get("name_style", {}))

    elif seg_type == "contact":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        body = styles.get("body_style", {})
        contact_style = {**body, "font_size_pt": body.get("font_size_pt", 10) - 1}
        _apply_run_style(run, contact_style)

    elif seg_type == "header":
        p = doc.add_paragraph()
        run = p.add_run(text.upper() if styles.get("section_header_style", {}).get("all_caps") else text)
        _apply_run_style(run, styles.get("section_header_style", {}))
        if styles.get("section_header_style", {}).get("bottom_border"):
            _add_bottom_border(p)
        pf = p.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(2)

    elif seg_type == "body":
        p = doc.add_paragraph()
        run = p.add_run(text)
        _apply_run_style(run, styles.get("body_style", {}))
        body_style = styles.get("body_style", {})
        if body_style.get("line_spacing_pt"):
            p.paragraph_format.line_spacing = Pt(body_style["line_spacing_pt"])

    elif seg_type == "bullet":
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        _apply_run_style(run, styles.get("body_style", {}))
        bullet_style = styles.get("bullet_style", {})
        pf = p.paragraph_format
        if bullet_style.get("indent_left_dxa"):
            pf.left_indent = Twips(bullet_style["indent_left_dxa"])
        if bullet_style.get("space_after_pt"):
            pf.space_after = Pt(bullet_style["space_after_pt"])


def render_docx(
    resume_text: str, styles: dict, template_docx_bytes: bytes = None
) -> bytes:
    segments = parse_tagged_resume(resume_text)

    if template_docx_bytes:
        doc = Document(BytesIO(template_docx_bytes))
        # Clear existing paragraph content but preserve styles
        for p in doc.paragraphs:
            p.clear()
        # Remove cleared empty paragraphs
        body = doc.element.body
        for p_elem in body.findall(qn("w:p")):
            body.remove(p_elem)
    else:
        doc = Document()

    # Apply page dimensions
    page = styles.get("page", {})
    if doc.sections:
        section = doc.sections[0]
        if page.get("width_dxa"):
            section.page_width = Twips(page["width_dxa"])
        if page.get("height_dxa"):
            section.page_height = Twips(page["height_dxa"])
        _set_page_margins(section, page)

    for segment in segments:
        _add_segment_to_doc(doc, segment, styles)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _dxa_to_pt(dxa: int) -> float:
    return dxa / 20.0


def render_pdf(resume_text: str, styles: dict) -> bytes:
    segments = parse_tagged_resume(resume_text)
    page = styles.get("page", {})
    name_s = styles.get("name_style", {})
    header_s = styles.get("section_header_style", {})
    body_s = styles.get("body_style", {})
    bullet_s = styles.get("bullet_style", {})

    page_w = _dxa_to_pt(page.get("width_dxa", 12240))
    page_h = _dxa_to_pt(page.get("height_dxa", 15840))
    m_top = _dxa_to_pt(page.get("margin_top_dxa", 1080))
    m_bot = _dxa_to_pt(page.get("margin_bottom_dxa", 1080))
    m_left = _dxa_to_pt(page.get("margin_left_dxa", 1080))
    m_right = _dxa_to_pt(page.get("margin_right_dxa", 1080))

    css = f"""
    @page {{
        size: {page_w}pt {page_h}pt;
        margin: {m_top}pt {m_right}pt {m_bot}pt {m_left}pt;
    }}
    body {{
        font-family: '{body_s.get("font_name", "Calibri")}', Calibri, Arial, sans-serif;
        font-size: {body_s.get("font_size_pt", 10)}pt;
        color: #{body_s.get("color_hex", "000000")};
        line-height: {body_s.get("line_spacing_pt", 12) / max(body_s.get("font_size_pt", 10), 1)};
        margin: 0;
        padding: 0;
    }}
    .name {{
        font-family: '{name_s.get("font_name", "Calibri")}', Calibri, Arial, sans-serif;
        font-size: {name_s.get("font_size_pt", 18)}pt;
        font-weight: {"bold" if name_s.get("bold") else "normal"};
        color: #{name_s.get("color_hex", "1a2744")};
        text-align: center;
        margin-bottom: 2pt;
    }}
    .contact {{
        font-size: {max(body_s.get("font_size_pt", 10) - 1, 8)}pt;
        text-align: center;
        color: #{body_s.get("color_hex", "000000")};
        margin-bottom: 6pt;
    }}
    .header {{
        font-family: '{header_s.get("font_name", "Calibri")}', Calibri, Arial, sans-serif;
        font-size: {header_s.get("font_size_pt", 11)}pt;
        font-weight: {"bold" if header_s.get("bold") else "normal"};
        color: #{header_s.get("color_hex", "1a2744")};
        {"text-transform: uppercase;" if header_s.get("all_caps") else ""}
        {"border-bottom: 1px solid #" + header_s.get("color_hex", "1a2744") + ";" if header_s.get("bottom_border") else ""}
        padding-bottom: 2pt;
        margin-top: 8pt;
        margin-bottom: 3pt;
    }}
    .body {{
        margin-bottom: 2pt;
    }}
    ul {{
        margin: 0;
        padding-left: {_dxa_to_pt(bullet_s.get("indent_left_dxa", 360))}pt;
    }}
    li {{
        margin-bottom: {bullet_s.get("space_after_pt", 2)}pt;
    }}
    """

    html_parts = [
        "<!DOCTYPE html><html><head>",
        f"<style>{css}</style>",
        "</head><body>",
    ]

    in_bullet_list = False

    for segment in segments:
        if segment["type"] == "bullet":
            if not in_bullet_list:
                html_parts.append("<ul>")
                in_bullet_list = True
            html_parts.append(f'<li>{_escape_html(segment["text"])}</li>')
        else:
            if in_bullet_list:
                html_parts.append("</ul>")
                in_bullet_list = False

            t = _escape_html(segment["text"])
            if segment["type"] == "name":
                html_parts.append(f'<div class="name">{t}</div>')
            elif segment["type"] == "contact":
                html_parts.append(f'<div class="contact">{t}</div>')
            elif segment["type"] == "header":
                html_parts.append(f'<div class="header">{t}</div>')
            elif segment["type"] == "body":
                html_parts.append(f'<div class="body">{t}</div>')

    if in_bullet_list:
        html_parts.append("</ul>")

    html_parts.append("</body></html>")
    html_str = "\n".join(html_parts)

    from weasyprint import HTML
    return HTML(string=html_str).write_pdf()


def _escape_html(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )
